"""Monta o artigo completo em Word a partir de artigo_metodologia.md,
artigo_resultados.md e references.md — mais Introducao e Conclusoes novas
(redigidas nesta etapa, citando so material ja verificado).

Uso: python -m scripts.gerar_artigo_word
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from scripts.word_utils import add_inline_markdown, convert_markdown_body

OUT_PATH = Path("artigo_design_inibidores.docx")

FIGURE_AFTER_HEADING = {
    "3.11f Réplicas Reais de MD (n=3) — Reprodutibilidade da Estabilidade (2026-07-19)":
        "outputs/figuras_artigo/fig1_replicas_md.png",
    "3.11g Persistência Competitiva Real — Ocupância do Bolso S1 ao Longo da Trajetória (2026-07-19)":
        "outputs/figuras_artigo/fig2_ocupancia_s1.png",
    "3.11i Assinatura Digital de Interação — Síntese das Frentes 2+3 (2026-07-19)":
        "outputs/figuras_artigo/fig4_fingerprint.png",
    "3.11j Fechamento do R3 — Matriz Consolidada TOP-13 × 11 Espécies (2026-07-19)":
        "outputs/figuras_artigo/fig3_cross_species.png",
}

INTRODUCAO = """
Lepidópteros-praga como *Spodoptera frugiperda* e *Anticarsia gemmatalis* causam perdas
significativas em culturas de importância econômica, e a resistência crescente a toxinas Bt e a
inseticidas químicos tem motivado a busca por estratégias complementares de controle [3][14][20].
Inibidores de proteases digestivas — em particular de tripsinas, enzimas centrais na digestão
proteica de larvas de Lepidoptera — representam uma dessas estratégias, atuando diretamente sobre
a fisiologia digestiva do inseto [1][27].

Inibidores naturais do tipo Kunitz/BPTI já demonstraram atividade real contra tripsinas de
Lepidoptera, incluindo inibição não-competitiva tight-binding com efeito mensurável sobre
sobrevivência larval [2] e sinergia com toxinas Bt [8]. No entanto, esses inibidores enfrentam
limitações práticas recorrentes: suscetibilidade à própria clivagem proteolítica quando contêm
sítios internos reconhecíveis pela protease-alvo, o que motiva estratégias de proteção química
como substituição por D-aminoácidos e ciclização [9][10][11].

Ferramentas de IA generativa para design de proteínas — RFdiffusion, para geração de novos
esqueletos estruturais [17], e ProteinMPNN, para desenho de sequência sobre um esqueleto dado
[18], além de extensões recentes para macrociclos [19] — abrem uma via alternativa à triagem de
inibidores naturais: gerar candidatos peptídicos inéditos, ancorados computacionalmente ao sítio
catalítico do alvo, e validá-los com o mesmo rigor experimental in silico já estabelecido para
docking molecular [23], dinâmica molecular [24] e refinamento de interface [25][26].

Este trabalho descreve um pipeline computacional multiagente que integra essas ferramentas para
o design racional de novos peptídeos inibidores de tripsinas de Lepidoptera, com ênfase explícita
em validação real e reprodutível em cada etapa — não apenas escore de docking estático, mas
estabilidade em réplicas independentes de dinâmica molecular, persistência real de contato com o
sítio catalítico, especificidade frente a organismos não-alvo, e amplo espectro contra múltiplas
espécies-praga de Lepidoptera.
""".strip()

CONCLUSOES = """
O pipeline multiagente gerou e avaliou, com dado real em cada etapa, um espaço amplo de candidatos
a inibidor peptídico de tripsina: 330 esqueletos reais via RFdiffusion, mais de 120 mil sequências
via ProteinMPNN, e milhares de dockings reais contra o receptor primário e 11 espécies de
Lepidoptera-praga. Desse espaço, um subconjunto de 13 candidatos (TOP-13) foi submetido a teste
aprofundado — réplicas reais de dinâmica molecular, persistência de contato com o sítio catalítico,
tríade catalítica via PLIP, e matriz cross-species completa.

O achado metodológico mais importante deste trabalho é que a estabilidade avaliada por uma única
réplica de dinâmica molecular não é confiável isoladamente: os dois candidatos historicamente
classificados como "mais estáveis do pipeline" a partir de réplica única
(RLREELKKAEEWLEKRRKEE, 0,294 nm; VRTRR, 0,194 nm) mostraram desvio-padrão real de 0,606 nm e
0,360 nm quando testados em réplicas independentes — variação grande o suficiente para invalidar a
classificação original. Apenas quatro candidatos (SRTRR, VRYRR, VRRPR, HRPRRPR) confirmaram
estabilidade reprodutível (DP<0,05 nm) nas três réplicas reais.

A análise de persistência real do contato com o bolso catalítico S1 refinou ainda mais esse
quadro: `VRRPR`, apesar de estabilidade reprodutível por RMSD do complexo inteiro, na verdade perde
contato real com o sítio ao longo das réplicas (ocupância a 6 Å caindo de 67,6% para 0,15% entre
réplicas) — evidência de que baixa variância estrutural não implica permanência funcional no sítio
catalítico. `VRYRR` destacou-se como o único candidato com contato tipo salt-bridge real e
constante mesmo no corte mais estrito (4 Å) nas três réplicas.

Dois bloqueadores permanecem sem solução simultânea em nenhum candidato: especificidade real
(0/23 candidatos aprovados com margem de seletividade ≥2,0 kcal/mol frente a tripsina humana e
*Apis mellifera*) e resistência proteolítica universal entre os candidatos de maior afinidade
(susceptibilidade generalizada a autoclivagem por resíduos P1-internos básicos). Um candidato
(SEEEVLAANEAYAAAHTAYN) reuniu pela primeira vez resistência estrutural real (ausência de sítios
K/R internos), afinidade competitiva e estabilidade em MD — mas ainda sem margem de especificidade
comprovada.

A verificação cruzada e independente de cada etapa computacional — incluindo a identificação e
correção de um erro real no algoritmo de detecção de tríade catalítica para uma das espécies
testadas, e a confirmação de que as demais tríades (incluindo as dos quatro receptores primários)
correspondem de fato ao mínimo geométrico global entre resíduos His/Asp/Ser — foi tratada como
parte integrante do método, não como etapa opcional. Os candidatos reprodutivelmente estáveis e
com contato real confirmado (`SRTRR`, `VRYRR`, `VRRPR`, `HRPRRPR`) constituem a base mais sólida
para as próximas etapas do projeto, ainda pendentes de resolução de especificidade e resistência
proteolítica antes de qualquer indicação para síntese e validação experimental.
""".strip()


def add_references(document, references_path: Path) -> None:
    text = references_path.read_text(encoding="utf-8")
    entries = re.split(r"\n(?=\*\*\[\d+\]\*\*)", text.split("## I.", 1)[1] if "## I." in text else text)
    document.add_heading("Referências", level=1)
    for block in entries:
        block = block.strip()
        if not block or not block.startswith("**["):
            continue
        lines = [l for l in block.splitlines() if l.strip()]
        citation_lines = [l for l in lines if not l.strip().startswith(">")]
        p = document.add_paragraph()
        add_inline_markdown(p, " ".join(l.strip() for l in citation_lines))


def build_cover(document) -> None:
    title = document.add_heading(
        "Design Racional de Inibidores Peptídicos de Tripsinas de Lepidoptera por IA Generativa",
        level=0,
    )
    p = document.add_paragraph()
    p.add_run("Universidade Federal de Viçosa (UFV) — Viçosa-MG, Brasil").italic = True
    document.add_page_break()


def main():
    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    build_cover(document)

    document.add_heading("Introdução", level=1)
    for para in INTRODUCAO.split("\n\n"):
        p = document.add_paragraph()
        add_inline_markdown(p, para.replace("\n", " "))

    metodologia_lines = Path("artigo_metodologia.md").read_text(encoding="utf-8").splitlines()
    convert_markdown_body(document, metodologia_lines[2:])  # pula titulo (linha 1) + linha em branco

    resultados_all = Path("artigo_resultados.md").read_text(encoding="utf-8").splitlines()
    resultados_body = resultados_all[19:1206]  # linhas 20-1206 (1-indexed) = corpo real, sem changelog/refs embutidas
    convert_markdown_body(document, resultados_body, figure_after_heading=FIGURE_AFTER_HEADING)

    document.add_heading("Conclusões", level=1)
    for para in CONCLUSOES.split("\n\n"):
        p = document.add_paragraph()
        add_inline_markdown(p, para.replace("\n", " "))

    add_references(document, Path("references.md"))

    document.save(str(OUT_PATH))
    print(f"Salvo: {OUT_PATH.resolve()}")


if __name__ == "__main__":
    main()
