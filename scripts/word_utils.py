"""Conversao Markdown -> python-docx para o artigo_design_inibidores.docx.

Formatacao inline (negrito/italico/codigo) via markdown-it-py (tokenizer real,
nao regex proprio) para lidar corretamente com combinacoes dentro de celulas
de tabela e paragrafos. Estrutura de blocos (titulos/tabelas/paragrafos) e
percorrida linha-a-linha, dado o formato real e consistente dos 3 arquivos
fonte (ver plano).
"""
from markdown_it import MarkdownIt

_md_inline = MarkdownIt()


def add_inline_markdown(paragraph, text: str) -> None:
    """Adiciona texto com negrito/italico/codigo real (via tokens
    markdown-it-py) como runs no paragrafo python-docx."""
    tokens = _md_inline.parseInline(text)[0].children
    bold_depth = 0
    italic_depth = 0
    for tok in tokens:
        if tok.type == "strong_open":
            bold_depth += 1
        elif tok.type == "strong_close":
            bold_depth -= 1
        elif tok.type == "em_open":
            italic_depth += 1
        elif tok.type == "em_close":
            italic_depth -= 1
        elif tok.type == "code_inline":
            run = paragraph.add_run(tok.content)
            run.font.name = "Consolas"
            run.bold = (bold_depth > 0) or None
            run.italic = (italic_depth > 0) or None
        elif tok.type == "text" and tok.content:
            run = paragraph.add_run(tok.content)
            run.bold = (bold_depth > 0) or None
            run.italic = (italic_depth > 0) or None


def parse_markdown_table(lines: list) -> tuple:
    """Parseia um bloco de tabela Markdown real (header, separador '---',
    linhas de dado) em (headers, rows) de strings cruas (formatacao inline
    aplicada depois, na insercao no docx)."""
    def split_row(line: str) -> list:
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    headers = split_row(lines[0])
    rows = [split_row(line) for line in lines[2:] if line.strip()]
    return headers, rows


def add_markdown_table(document, headers: list, rows: list) -> None:
    """Cria uma tabela real no docx, com negrito/italico/codigo reais em
    cada celula (via add_inline_markdown) — nao so texto plano."""
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.paragraphs[0].clear()
        add_inline_markdown(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.paragraphs[0].clear()
            add_inline_markdown(cell.paragraphs[0], val)


from docx.shared import Inches


import re

_BULLET_RE = re.compile(r"^-\s+(.*)")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)")


def convert_markdown_body(document, lines: list, figure_after_heading: dict = None) -> None:
    """Percorre linhas de um corpo Markdown real (##/### cabecalhos,
    paragrafos, listas com marcador '- '/'N. ' (item por linha, sem quebra em
    continuacao — formato real confirmado em artigo_resultados.md), tabelas
    com legenda **Tabela N.**, '---' como separador ignorado) e monta o docx.
    figure_after_heading insere uma imagem logo apos o primeiro paragrafo de
    texto seguinte ao cabecalho indicado."""
    figure_after_heading = figure_after_heading or {}
    pending_figure = None
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            pending_figure = figure_after_heading.get(stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            p = document.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet_match.group(1))
            i += 1
            continue

        numbered_match = _NUMBERED_RE.match(stripped)
        if numbered_match:
            p = document.add_paragraph(style="List Number")
            add_inline_markdown(p, numbered_match.group(1))
            i += 1
            continue

        if stripped.startswith("**Tabela") or stripped.startswith("**Table"):
            caption_lines = [stripped]
            i += 1
            # caption pode continuar em linhas adicionais (ate blank line ou tabela)
            while i < n and lines[i].strip() and not lines[i].strip().startswith("|"):
                caption_lines.append(lines[i].strip())
                i += 1
            caption_p = document.add_paragraph()
            add_inline_markdown(caption_p, " ".join(caption_lines))
            # proxima linha nao-vazia deve ser a tabela
            while i < n and not lines[i].strip():
                i += 1
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if table_lines:  # so processa se houver tabela
                headers, rows = parse_markdown_table(table_lines)
                add_markdown_table(document, headers, rows)
            continue

        # paragrafo normal: junta linhas contiguas nao-vazias, nao-tabela,
        # nao-cabecalho, nao-lista (defensivo — no formato real observado,
        # listas sempre tem linha em branco antes, mas nao custa garantir)
        para_lines = [stripped]
        i += 1
        while (i < n and lines[i].strip()
               and not lines[i].strip().startswith(("#", "|", "**Tabela", "**Table", "---"))
               and not _BULLET_RE.match(lines[i].strip())
               and not _NUMBERED_RE.match(lines[i].strip())):
            para_lines.append(lines[i].strip())
            i += 1
        p = document.add_paragraph()
        add_inline_markdown(p, " ".join(para_lines))

        if pending_figure:
            document.add_picture(pending_figure, width=Inches(6.0))
            pending_figure = None
