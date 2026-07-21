from docx import Document
from scripts.word_utils import (
    add_inline_markdown,
    parse_markdown_table,
    add_markdown_table,
    convert_markdown_body,
)

def test_add_inline_markdown_plain_text():
    doc = Document()
    p = doc.add_paragraph()
    add_inline_markdown(p, "texto simples")
    assert len(p.runs) == 1
    assert p.runs[0].text == "texto simples"
    assert p.runs[0].bold is None
    assert p.runs[0].italic is None

def test_add_inline_markdown_bold_and_italic():
    doc = Document()
    p = doc.add_paragraph()
    add_inline_markdown(p, "**negrito** e *itálico* e normal")
    texts = [(r.text, bool(r.bold), bool(r.italic)) for r in p.runs]
    assert texts == [
        ("negrito", True, False),
        (" e ", False, False),
        ("itálico", False, True),
        (" e normal", False, False),
    ]

def test_add_inline_markdown_code_inline():
    doc = Document()
    p = doc.add_paragraph()
    add_inline_markdown(p, "roda `scripts/prep_pdbs.py` real")
    assert p.runs[1].text == "scripts/prep_pdbs.py"
    assert p.runs[1].font.name == "Consolas"

def test_add_inline_markdown_code_inline_inherits_bold():
    doc = Document()
    p = doc.add_paragraph()
    add_inline_markdown(p, "**Bug (commit `e9024bc`):**")
    code_run = next(r for r in p.runs if r.text == "e9024bc")
    assert code_run.bold is True
    assert code_run.font.name == "Consolas"

def test_add_inline_markdown_same_type_nested_emphasis():
    doc = Document()
    p = doc.add_paragraph()
    add_inline_markdown(p, "**outer __inner__ outer**")
    texts = [(r.text, bool(r.bold)) for r in p.runs]
    assert texts == [
        ("outer ", True),
        ("inner", True),
        (" outer", True),
    ]

def test_parse_markdown_table_real_format():
    lines = [
        "| Modelo       | Res. catalítico 1 | Asp (tríade) |",
        "|--------------|-------------------|--------------|",
        "| ACR157       | His69             | Asp114       |",
        "| **Consenso** | —                 | —            |",
    ]
    headers, rows = parse_markdown_table(lines)
    assert headers == ["Modelo", "Res. catalítico 1", "Asp (tríade)"]
    assert rows == [
        ["ACR157", "His69", "Asp114"],
        ["**Consenso**", "—", "—"],
    ]

def test_add_markdown_table_creates_real_docx_table():
    doc = Document()
    headers = ["A", "B"]
    rows = [["1", "2"], ["3", "**4**"]]
    add_markdown_table(doc, headers, rows)
    table = doc.tables[0]
    assert len(table.rows) == 3  # header + 2 data rows
    assert table.cell(0, 0).paragraphs[0].runs[0].text == "A"
    assert table.cell(0, 0).paragraphs[0].runs[0].bold is True
    assert table.cell(2, 1).paragraphs[0].runs[0].text == "4"
    assert table.cell(2, 1).paragraphs[0].runs[0].bold is True

def test_convert_markdown_body_headings_paragraphs_tables():
    lines = [
        "### 3.1 Título da Seção",
        "",
        "Um parágrafo com **negrito** real.",
        "",
        "**Tabela 1.** Legenda.",
        "",
        "| A | B |",
        "|---|---|",
        "| 1 | 2 |",
        "",
        "---",
        "",
        "### 3.2 Outra Seção",
        "",
        "Outro parágrafo.",
    ]
    doc = Document()
    convert_markdown_body(doc, lines)

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert headings == ["3.1 Título da Seção", "3.2 Outra Seção"]
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(1, 0).paragraphs[0].runs[0].text == "1"

def test_convert_markdown_body_handles_consecutive_list_items():
    # Formato real confirmado em artigo_resultados.md:143-145 — itens de lista
    # consecutivos, SEM linha em branco entre eles, cada um autocontido numa linha.
    lines = [
        "### Seção com lista",
        "",
        "- **Item A**: descrição do item A.",
        "- **Item B**: descrição do item B.",
        "- **Item C**: descrição do item C.",
        "",
        "1. Primeiro item numerado.",
        "2. Segundo item numerado.",
    ]
    doc = Document()
    convert_markdown_body(doc, lines)
    body_paragraphs = [p for p in doc.paragraphs if not p.style.name.startswith("Heading")]
    # 3 itens de bullet + 2 itens numerados = 5 paragrafos distintos (nao 2 parágrafos mesclados)
    assert len(body_paragraphs) == 5
    assert body_paragraphs[0].runs[0].text == "Item A"  # negrito preservado, marcador removido
    assert body_paragraphs[0].style.name == "List Bullet"
    assert body_paragraphs[3].style.name == "List Number"
    assert "Primeiro item numerado." in body_paragraphs[3].text

def test_convert_markdown_body_inserts_figure_after_heading():
    lines = [
        "### 3.11f Réplicas Reais de MD",
        "",
        "Parágrafo real da seção.",
        "",
        "### 3.11g Próxima Seção",
    ]
    doc = Document()
    convert_markdown_body(
        doc, lines,
        figure_after_heading={"3.11f Réplicas Reais de MD": "outputs/figuras_artigo/fig1_replicas_md.png"},
    )
    has_image = any(
        run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        for p in doc.paragraphs for run in p.runs
    )
    assert has_image
